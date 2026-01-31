# app.py
from flask import Flask, request, jsonify, Response, send_file, send_from_directory
from flask_cors import CORS
import pymupdf4llm
import pymupdf
import os
import logging
import traceback
import json
import time
from datetime import datetime
from threading import Thread
import uuid
import sys
from io import StringIO, BytesIO
import re
import pypandoc
import traceback 
import subprocess
import signal
import tomllib
from docx import Document
from docx.shared import Pt, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def _get_static_folder():
    env_static = os.environ.get("PDF3MD_STATIC_DIR")
    if env_static and os.path.exists(env_static):
        return env_static

    if getattr(sys, 'frozen', False) and hasattr(sys, '_MEIPASS'):
        base_dir = sys._MEIPASS
    else:
        base_dir = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(base_dir, "dist")

app = Flask(__name__, static_folder=_get_static_folder(), static_url_path="")

# Default CORS origins for local development
default_origins = [
    "http://localhost:5173",    # Vite dev server default
    "http://localhost:3000",    # Production frontend port in Docker
    "http://127.0.0.1:5173",
    "http://127.0.0.1:3000",
    "*"  # Allow all origins for local hosting (permissive for local use)
]

# Get additional origins from environment variable
additional_origins_str = os.environ.get('ALLOWED_CORS_ORIGINS')
all_allowed_origins = list(default_origins) # Start with a copy of defaults

if additional_origins_str:
    # Split the comma-separated string and strip whitespace from each origin
    custom_origins = [origin.strip() for origin in additional_origins_str.split(',')]
    
    # For each base domain, add variations with common ports
    expanded_origins = []
    for origin in custom_origins:
        expanded_origins.append(origin)  # Add the base domain
        
        # Extract the domain without protocol
        domain_match = re.match(r'(https?://)?(.*)', origin)
        if domain_match:
            protocol = domain_match.group(1) or 'http://'  # Default to http:// if no protocol
            domain = domain_match.group(2)
            
            # Add common port variations
            expanded_origins.append(f"{protocol}{domain}:3000")  # Production frontend
            expanded_origins.append(f"{protocol}{domain}:5173")  # Dev frontend
            expanded_origins.append(f"{protocol}{domain}:6201")  # Backend
    
    all_allowed_origins.extend(expanded_origins)

# Remove duplicates by converting to a set and back to a list
final_origins = list(set(all_allowed_origins))

# Set up logging (ensure logger is configured before use, especially for the CORS log line)
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

logger.info(f"Initializing CORS with origins: {final_origins}") # Log the origins

# Apply CORS with very permissive settings for local hosting
CORS(
    app,
    origins="*",  # Allow all origins
    methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"],
    allow_headers=["Content-Type", "Authorization", "X-Requested-With", "Origin", "Accept"],
    supports_credentials=False
)

# Add manual CORS headers as backup
@app.after_request
def after_request(response):
    response.headers.add('Access-Control-Allow-Origin', '*')
    response.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization,X-Requested-With,Origin,Accept')
    response.headers.add('Access-Control-Allow-Methods', 'GET,PUT,POST,DELETE,OPTIONS')
    response.headers.add('Access-Control-Expose-Headers', 'Content-Disposition')
    return response

# Store conversion progress
conversion_progress = {}

def _load_version_meta():
    version = "0.0.0"
    release_date = "unknown"
    developer = "unknown"

    project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    pyproject_path = os.path.join(project_root, "pyproject.toml")
    if os.path.exists(pyproject_path):
        try:
            with open(pyproject_path, "rb") as f:
                data = tomllib.load(f)
            version = data.get("project", {}).get("version", version)
            tool = data.get("tool", {}).get("pdf3md", {})
            release_date = tool.get("release_date", release_date)
            developer = tool.get("developer", developer)
            return version, release_date, developer
        except Exception:
            pass

    build_meta = os.path.join(os.path.dirname(__file__), "build_meta.json")
    if os.path.exists(build_meta):
        try:
            with open(build_meta, "r", encoding="utf-8") as f:
                payload = json.load(f)
            version = payload.get("version", version)
            release_date = payload.get("release_date", release_date)
            developer = payload.get("developer", developer)
            return version, release_date, developer
        except Exception:
            pass

    version_json = os.path.join(os.path.dirname(__file__), "version.json")
    if os.path.exists(version_json):
        try:
            with open(version_json, "r", encoding="utf-8") as f:
                payload = json.load(f)
            version = payload.get("version", version)
            release_date = payload.get("release_date", release_date)
            developer = payload.get("developer", developer)
        except Exception:
            pass

    return version, release_date, developer

def _get_git_info():
    try:
        commit = subprocess.check_output(["git", "rev-parse", "--short", "HEAD"], text=True).strip()
        branch = subprocess.check_output(["git", "rev-parse", "--abbrev-ref", "HEAD"], text=True).strip()
        dirty = subprocess.check_output(["git", "status", "--porcelain"], text=True).strip() != ""
        describe = subprocess.check_output(["git", "describe", "--tags", "--always"], text=True).strip()
        return commit, branch, dirty, describe
    except Exception:
        return None, None, None, None

class ProgressCapture:
    """Capture progress output from pymupdf4llm"""
    def __init__(self, conversion_id, total_pages):
        self.conversion_id = conversion_id
        self.total_pages = total_pages
        self.original_stdout = sys.stdout
        self.original_stderr = sys.stderr
        
    def __enter__(self):
        sys.stdout = self
        sys.stderr = self
        return self
        
    def __exit__(self, exc_type, exc_val, exc_tb):
        sys.stdout = self.original_stdout
        sys.stderr = self.original_stderr
        
    def write(self, text):
        # Write to original stdout/stderr
        self.original_stdout.write(text)
        self.original_stdout.flush()
        
        # Parse progress from pymupdf4llm output
        if self.conversion_id in conversion_progress:
            # Look for progress patterns like "[====    ] (5/26)" or "Processing page 5 of 26"
            progress_match = re.search(r'\(\s*(\d+)/(\d+)\)', text)
            if progress_match:
                current_page = int(progress_match.group(1))
                total_pages = int(progress_match.group(2))
                
                # Calculate progress percentage (reserve 10% for finalization)
                progress_percent = int((current_page / total_pages) * 85) + 10
                
                conversion_progress[self.conversion_id].update({
                    'progress': progress_percent,
                    'stage': f'Processing page {current_page} of {total_pages}...',
                    'current_page': current_page,
                    'total_pages': total_pages
                })
                
    def flush(self):
        self.original_stdout.flush()

def convert_pdf_with_progress(temp_path, conversion_id, filename):
    """Convert PDF with real progress tracking"""
    try:
        # Get PDF metadata first
        doc = pymupdf.open(temp_path)
        total_pages = len(doc)
        file_size = os.path.getsize(temp_path)
        doc.close()
        
        # Update progress: Starting conversion
        conversion_progress[conversion_id] = {
            'progress': 0,
            'stage': 'Starting conversion...',
            'total_pages': total_pages,
            'current_page': 0,
            'filename': filename,
            'file_size': file_size,
            'status': 'processing'
        }
        
        # Start conversion with progress tracking
        logger.info(f'Starting PDF conversion for {filename}')
        
        # Update progress: Converting to markdown
        conversion_progress[conversion_id].update({
            'progress': 5,
            'stage': 'Initializing conversion...'
        })
        
        # Capture progress output from pymupdf4llm
        with ProgressCapture(conversion_id, total_pages):
            # Actual conversion - this is where the real work happens
            markdown = pymupdf4llm.to_markdown(temp_path)
        
        # Update progress: Finalizing
        conversion_progress[conversion_id].update({
            'progress': 95,
            'stage': 'Finalizing conversion...'
        })
        
        time.sleep(0.5)  # Brief pause for finalization
        
        # Format file size
        def format_file_size(size_bytes):
            if size_bytes < 1024:
                return f"{size_bytes} B"
            elif size_bytes < 1024 * 1024:
                return f"{size_bytes / 1024:.1f} KB"
            else:
                return f"{size_bytes / (1024 * 1024):.1f} MB"
        
        # Complete conversion
        result = {
            'markdown': markdown,
            'filename': filename,
            'fileSize': format_file_size(file_size),
            'pageCount': total_pages,
            'timestamp': datetime.now().isoformat(),
            'success': True
        }
        
        conversion_progress[conversion_id].update({
            'progress': 100,
            'stage': 'Conversion complete!',
            'status': 'completed',
            'result': result
        })
        
        logger.info('Conversion successful')
        
    except Exception as e:
        logger.error(f'Conversion error: {str(e)}')
        logger.error(traceback.format_exc())
        conversion_progress[conversion_id] = {
            'progress': 0,
            'stage': f'Error: {str(e)}',
            'status': 'error',
            'error': str(e)
        }

@app.route('/convert', methods=['POST'])
def convert():
    try:
        # --- BEGIN ADDED CLEANUP ---
        # Proactively clean up any orphaned temp_*.pdf files
        # Assumes temp files are in the same directory as app.py
        # The current working directory for app.py when run via Docker is /app/pdf3md
        # but when run locally for dev, it's where app.py is.
        # os.abspath('.') will give the correct directory in both cases if app.py is the entrypoint.
        current_dir = os.path.abspath(os.path.dirname(__file__)) # More robust way to get script's dir
        logger.info(f"Checking for orphaned temp files in: {current_dir}")
        cleaned_count = 0
        for filename in os.listdir(current_dir):
            if filename.startswith('temp_') and filename.endswith('.pdf'):
                # Further check if it's an orphaned file (not in current conversion_progress)
                # This check is a bit tricky because conversion_id is generated *after* this cleanup.
                # For simplicity, we'll clean up any file matching the pattern.
                # A more sophisticated check might involve checking if the conversion_id part of the filename
                # corresponds to an active or very recent conversion.
                # However, given the problem is orphaned files, a broad cleanup is likely fine.
                file_path_to_delete = os.path.join(current_dir, filename)
                try:
                    os.remove(file_path_to_delete)
                    logger.info(f"Proactively removed orphaned temp file: {file_path_to_delete}")
                    cleaned_count += 1
                except Exception as e_clean:
                    logger.error(f"Error removing orphaned temp file {file_path_to_delete}: {e_clean}")
        if cleaned_count > 0:
            logger.info(f"Proactively cleaned up {cleaned_count} orphaned temp PDF files.")
        # --- END ADDED CLEANUP ---

        if 'pdf' not in request.files:
            logger.error('No file in request')
            return jsonify({'error': 'No file uploaded'}), 400
            
        file = request.files['pdf']
        if file.filename == '':
            logger.error('Empty filename')
            return jsonify({'error': 'No file selected'}), 400

        # Generate unique conversion ID
        conversion_id = str(uuid.uuid4())
        
        # Save uploaded file temporarily
        temp_path = os.path.abspath(f'temp_{conversion_id}.pdf')
        logger.info(f'Saving file to {temp_path}')
        file.save(temp_path)
        
        # Start conversion in background thread
        thread = Thread(target=convert_pdf_with_progress, args=(temp_path, conversion_id, file.filename))
        thread.start()
        
        # Return conversion ID for progress tracking
        return jsonify({
            'conversion_id': conversion_id,
            'message': 'Conversion started',
            'success': True
        })
        
    except Exception as e:
        logger.error(f'Server error: {str(e)}')
        logger.error(traceback.format_exc())
        return jsonify({'error': f'Server error: {str(e)}', 'success': False}), 500

@app.route('/progress/<conversion_id>', methods=['GET'])
def get_progress(conversion_id):
    """Get conversion progress for a specific conversion ID"""
    try:
        if conversion_id not in conversion_progress:
            return jsonify({'error': 'Conversion not found'}), 404
        
        progress_data = conversion_progress[conversion_id].copy()
        
        # Clean up completed or errored conversions after sending response
        if progress_data.get('status') in ['completed', 'error']:
            # Clean up temp file
            temp_path = os.path.abspath(f'temp_{conversion_id}.pdf')
            if os.path.exists(temp_path):
                os.remove(temp_path)
                logger.info(f'Temp file removed: {temp_path}')
            
            # Remove from progress tracking after a delay to allow final fetch
            def cleanup_progress():
                time.sleep(5)  # Wait 5 seconds before cleanup
                if conversion_id in conversion_progress:
                    del conversion_progress[conversion_id]
            
            Thread(target=cleanup_progress).start()
        
        return jsonify(progress_data)
        
    except Exception as e:
        logger.error(f'Progress error: {str(e)}')
        return jsonify({'error': f'Progress error: {str(e)}'}), 500

def markdown_to_docx(markdown_text, filename="document"):
    """Convert markdown text to a Word document using Pandoc."""
    temp_docx_path = None
    try:
        _ensure_pandoc_available()
        # pypandoc.ensure_pandoc_installed() # Optional

        # Generate a unique temporary filename for the DOCX output
        temp_docx_filename = f"temp_pandoc_output_{uuid.uuid4()}.docx"
        # Ensure the temp path is absolute, similar to how temp PDFs are handled
        temp_docx_path = os.path.abspath(temp_docx_filename)

        logger.debug(f"Attempting to convert markdown to docx for filename: {filename} using pandoc, outputting to {temp_docx_path}")
        
        pypandoc.convert_text(
            markdown_text,
            'docx',
            format='md',
            outputfile=temp_docx_path
            # extra_args=['--verbose'] # Uncomment for pandoc verbose logging if needed
        )

        try:
            _apply_docx_formatting(temp_docx_path)
        except Exception as format_error:
            logger.warning(f"Post-processing DOCX formatting failed: {format_error}")
        
        # Read the generated DOCX file into a BytesIO buffer
        with open(temp_docx_path, 'rb') as f:
            output_docx_bytes = f.read()
            
        doc_buffer = BytesIO(output_docx_bytes)
        doc_buffer.seek(0)
        
        logger.info(f"Successfully converted markdown to docx for {filename} using pandoc (via temp file: {temp_docx_path}).")
        return doc_buffer
        
    except FileNotFoundError: # Specifically catch if pandoc is not found
        logger.error('Pandoc not found. Please ensure Pandoc is installed and in your PATH.')
        logger.error(traceback.format_exc())
        raise RuntimeError('Pandoc not found. Conversion failed.') # Re-raise a more specific error
    except Exception as e:
        logger.error(f'Error converting markdown to docx using Pandoc: {str(e)}')
        logger.error(traceback.format_exc())
        raise e # Re-raise to be handled by the route
    finally:
        # Clean up the temporary DOCX file
        if temp_docx_path and os.path.exists(temp_docx_path):
            try:
                os.remove(temp_docx_path)
                logger.debug(f"Successfully removed temporary pandoc output file: {temp_docx_path}")
            except Exception as e_clean:
                logger.error(f"Error removing temporary pandoc output file {temp_docx_path}: {str(e_clean)}")

def _apply_docx_formatting(docx_path):
    doc = Document(docx_path)

    _apply_page_margins(doc)
    _remove_leading_metadata(doc)
    _remove_horizontal_rules(doc)
    _remove_shape_lines(doc)
    _add_page_numbers(doc)
    _apply_heading_sizes(doc)
    _format_tables(doc)

    doc.save(docx_path)

def _ensure_pandoc_available():
    if os.environ.get("PDF3MD_SKIP_PANDOC_DOWNLOAD", "0") == "1":
        return

    pandoc_env = os.environ.get("PYPANDOC_PANDOC")
    if pandoc_env and os.path.exists(pandoc_env):
        return

    if getattr(sys, 'frozen', False) and hasattr(sys, '_MEIPASS'):
        bundled = os.path.join(sys._MEIPASS, "pandoc", "pandoc")
        if os.path.exists(bundled):
            os.environ["PYPANDOC_PANDOC"] = bundled
            return

    app_support = os.path.join(os.path.expanduser("~"), "Library", "Application Support", "PDF3MD", "pandoc")
    pandoc_path = os.path.join(app_support, "pandoc")
    if os.path.exists(pandoc_path):
        os.environ["PYPANDOC_PANDOC"] = pandoc_path
        return

    os.makedirs(app_support, exist_ok=True)
    try:
        pypandoc.download_pandoc(targetfolder=app_support)
        if os.path.exists(pandoc_path):
            os.environ["PYPANDOC_PANDOC"] = pandoc_path
    except Exception as e:
        logger.error(f"Failed to download pandoc: {e}")

def _remove_leading_metadata(doc):
    paragraphs = list(doc.paragraphs)
    if not paragraphs:
        return

    to_delete = []
    for paragraph in paragraphs[:5]:
        text = paragraph.text.strip()
        if text.startswith("[[") and "#" in text:
            to_delete.append(paragraph)
            continue
        if text in {"---", "***", "___"}:
            to_delete.append(paragraph)
            continue
        if _has_horizontal_rule(paragraph):
            to_delete.append(paragraph)

    for paragraph in to_delete:
        _delete_paragraph(paragraph)

def _delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None

def _has_horizontal_rule(paragraph):
    p_pr = paragraph._element.find(qn("w:pPr"))
    if p_pr is None:
        return False
    p_borders = p_pr.find(qn("w:pBdr"))
    if p_borders is None:
        return False
    for edge in ("top", "bottom", "left", "right"):
        element = p_borders.find(qn(f"w:{edge}"))
        if element is not None and element.get(qn("w:val")):
            return True
    return False

def _remove_horizontal_rules(doc):
    for paragraph in list(doc.paragraphs):
        if _has_horizontal_rule(paragraph):
            _delete_paragraph(paragraph)

    for section in doc.sections:
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        for paragraph in list(section.header.paragraphs):
            if _has_horizontal_rule(paragraph):
                _delete_paragraph(paragraph)
            else:
                paragraph.clear()
        for paragraph in list(section.footer.paragraphs):
            if _has_horizontal_rule(paragraph):
                _delete_paragraph(paragraph)

def _remove_shape_lines(doc):
    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip():
            continue
        has_drawing = paragraph._element.findall(".//w:drawing", namespaces=paragraph._element.nsmap)
        has_pict = paragraph._element.findall(".//w:pict", namespaces=paragraph._element.nsmap)
        if has_drawing or has_pict:
            _delete_paragraph(paragraph)

def _apply_page_margins(doc):
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.3)
        section.bottom_margin = Inches(0.3)
        section.left_margin = Inches(0.79)
        section.right_margin = Inches(0.33)
        section.header_distance = Inches(0)
        section.footer_distance = Inches(0.2)

def _add_page_numbers(doc):
    for section in doc.sections:
        footer = section.footer
        if footer.paragraphs:
            paragraph = footer.paragraphs[0]
        else:
            paragraph = footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        run = paragraph.add_run()
        fld_char_begin = OxmlElement('w:fldChar')
        fld_char_begin.set(qn('w:fldCharType'), 'begin')

        instr_text = OxmlElement('w:instrText')
        instr_text.set(qn('xml:space'), 'preserve')
        instr_text.text = ' PAGE '

        fld_char_separate = OxmlElement('w:fldChar')
        fld_char_separate.set(qn('w:fldCharType'), 'separate')

        fld_char_end = OxmlElement('w:fldChar')
        fld_char_end.set(qn('w:fldCharType'), 'end')

        run._r.append(fld_char_begin)
        run._r.append(instr_text)
        run._r.append(fld_char_separate)
        run._r.append(fld_char_end)

def _apply_heading_sizes(doc):
    heading_sizes = {
        "Heading 1": 14,
        "Heading 2": 12,
        "Heading 3": 11,
        "Heading 4": 10,
        "Heading 5": 9,
        "Heading 6": 9
    }

    for style_name, size_pt in heading_sizes.items():
        if style_name in doc.styles:
            style = doc.styles[style_name]
            if style and style.font:
                style.font.size = Pt(size_pt)

    for paragraph in doc.paragraphs:
        if paragraph.style and paragraph.style.name in heading_sizes:
            size_pt = heading_sizes[paragraph.style.name]
            for run in paragraph.runs:
                run.font.size = Pt(size_pt)

def _format_tables(doc):
    if not doc.tables:
        return

    section = doc.sections[0]
    available_width = section.page_width - section.left_margin - section.right_margin

    for table in doc.tables:
        rows_count = len(table.rows)
        cols_count = len(table.columns)

        if cols_count >= 6:
            min_col_width = Inches(0.35)
            max_col_width = Inches(2.6)
            table_font_size = Pt(10)
        elif rows_count >= 6:
            min_col_width = Inches(0.45)
            max_col_width = Inches(3.0)
            table_font_size = Pt(10)
        elif cols_count == 2:
            min_col_width = Inches(0.5)
            max_col_width = Inches(3.2)
            table_font_size = Pt(11)
        else:
            min_col_width = Inches(0.45)
            max_col_width = Inches(3.0)
            table_font_size = Pt(11)

        table.style = "Table"
        table.autofit = False
        _set_table_borders(table)

        if table.rows:
            header_row = table.rows[0]
        else:
            header_row = None

        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                _set_cell_borders(cell)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = table_font_size
                        if header_row and row_index == 0:
                            run.bold = True
                    if header_row and row_index == 0:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if header_row and row_index == 0:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        if header_row:
            _normalize_header_labels(header_row)

        _align_table_columns(table)
        _adjust_table_column_widths(table, available_width, min_col_width, max_col_width, header_row)

def _adjust_table_column_widths(table, available_width, min_col_width, max_col_width, header_row):
    if not table.columns:
        return

    available_width_emu = int(available_width)
    min_col_width_emu = int(min_col_width)
    max_col_width_emu = int(max_col_width)

    header_texts = []
    if header_row:
        for cell in header_row.cells:
            header_texts.append(_cell_text(cell))
    else:
        header_texts = [""] * len(table.columns)

    column_max_lengths = []
    per_col_min = []
    for col_index, column in enumerate(table.columns):
        max_len = 1
        for cell in column.cells:
            cell_text = _cell_text(cell)
            if cell_text:
                max_len = max(max_len, len(cell_text))
        column_max_lengths.append(min(max_len, 120))
        per_col_min.append(min_col_width_emu)

    total_weight = sum(column_max_lengths) or 1
    weighted_lengths = []
    for idx, base in enumerate(column_max_lengths):
        header = header_texts[idx] if idx < len(header_texts) else ""
        weight = float(base)

        header_lower = header.lower()
        header_min_width = _header_min_width_in(header)
        if header_min_width:
            per_col_min[idx] = max(per_col_min[idx], int(Inches(header_min_width)))

        if "4-7" in header_lower and "data" in header_lower:
            weight *= 1.2
            if idx < len(per_col_min):
                per_col_min[idx] = int(Inches(0.95))
        elif any(key in header_lower for key in ["примеч", "назнач", "description", "comment"]):
            weight *= 1.6
        elif any(key in header_lower for key in ["устрой", "device"]):
            weight *= 1.2
        elif any(key in header_lower for key in ["байт", "byte", "cmd", "code", "op", "id", "№"]):
            weight *= 0.6
        elif "блок" in header_lower:
            weight *= 1.6
            if idx < len(per_col_min):
                per_col_min[idx] = int(Inches(0.9))

        column = table.columns[idx]
        texts = [_cell_text(cell) for cell in column.cells if _cell_text(cell)]
        if texts:
            numeric_like = sum(1 for t in texts if _is_numeric_like(t))
            if numeric_like / len(texts) >= 0.7:
                weight *= 0.6

        weighted_lengths.append(max(weight, 1))

    total_weight = sum(weighted_lengths) or 1
    raw_widths = [
        available_width_emu * (weight / total_weight)
        for weight in weighted_lengths
    ]

    clamped_widths = []
    for idx, width in enumerate(raw_widths):
        min_width = per_col_min[idx] if idx < len(per_col_min) else min_col_width_emu
        if width < min_width:
            clamped_widths.append(min_width)
        elif width > max_col_width_emu:
            clamped_widths.append(max_col_width_emu)
        else:
            clamped_widths.append(width)

    total_width = sum(clamped_widths) or available_width_emu
    scale = available_width_emu / total_width if total_width else 1
    final_widths = [int(width * scale) for width in clamped_widths]

    for col_index, width in enumerate(final_widths):
        table.columns[col_index].width = Emu(width)
        for cell in table.columns[col_index].cells:
            cell.width = Emu(width)

def _align_table_columns(table):
    if not table.columns:
        return

    for col_index, column in enumerate(table.columns):
        texts = []
        for cell in column.cells:
            text = _cell_text(cell)
            if text:
                texts.append(text)

        if not texts:
            continue

        numeric_like = sum(1 for t in texts if _is_numeric_like(t))
        avg_len = sum(len(t) for t in texts) / len(texts)
        should_center = (numeric_like / len(texts) >= 0.7) or avg_len <= 4

        if should_center:
            for cell in column.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def _is_numeric_like(text):
    value = text.strip()
    if not value:
        return False
    value = value.replace(",", ".")
    if re.fullmatch(r"[+-]?\d+(\.\d+)?", value):
        return True
    if re.fullmatch(r"0x[0-9a-fA-F]+", value):
        return True
    if re.fullmatch(r"[0-9a-fA-F]{2,}", value):
        return True
    return False

def _normalize_header_labels(header_row):
    for cell in header_row.cells:
        for paragraph in cell.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            if "Блок" in text:
                paragraph.clear()
                run = paragraph.add_run("Блок")
                run.font.size = Pt(10)
                run.bold = True
            if "Устройство" in text:
                paragraph.clear()
                run = paragraph.add_run("Устройство")
                run.font.size = Pt(10)
                run.bold = True
            if "1-й" in text and "байт" in text:
                paragraph.clear()
                run = paragraph.add_run("1-й байт")
                run.font.size = Pt(10)
                run.bold = True
            if "2-й" in text and "байт" in text:
                paragraph.clear()
                run = paragraph.add_run("2-й байт")
                run.font.size = Pt(10)
                run.bold = True
            if "4-7" in text and "Data" in text:
                paragraph.clear()
                run = paragraph.add_run("4-7 ?байт")
                run.font.size = Pt(10)
                run.bold = True
                run.add_break()
                run2 = paragraph.add_run("(Data)")
                run2.font.size = Pt(10)
                run2.bold = True

def _header_min_width_in(header_text):
    if not header_text:
        return None

    text_lower = header_text.lower()
    if "устрой" in text_lower or "device" in text_lower:
        return 1.2
    if "1-й" in text_lower and "байт" in text_lower:
        return 0.7
    if "2-й" in text_lower and "байт" in text_lower:
        return 0.7
    if "блок" in text_lower:
        return 0.7
    if "4-7" in text_lower and "data" in text_lower:
        return 0.95

    clean_len = len(re.sub(r"\s+", "", header_text))
    if clean_len <= 6:
        return 0.6
    if clean_len <= 12:
        return 0.85
    if clean_len <= 20:
        return 1.2
    return 1.4

def _set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.find(qn('w:tblBorders'))
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tbl_pr.append(borders)

    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        element = borders.find(qn(f'w:{edge}'))
        if element is None:
            element = OxmlElement(f'w:{edge}')
            borders.append(element)
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '8')
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), '000000')

def _set_cell_borders(cell):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.find(qn('w:tcBorders'))
    if borders is None:
        borders = OxmlElement('w:tcBorders')
        tc_pr.append(borders)

    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_tag = qn(f'w:{edge}')
        element = borders.find(edge_tag)
        if element is None:
            element = OxmlElement(f'w:{edge}')
            borders.append(element)
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '8')
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), '000000')

def _cell_text(cell):
    parts = []
    for paragraph in cell.paragraphs:
        if paragraph.text:
            parts.append(paragraph.text.strip())
    return " ".join(parts).strip()

@app.route('/convert-markdown-to-word', methods=['POST'])
def convert_markdown_to_word():
    """Convert markdown text to Word document"""
    try:
        data = request.get_json()
        
        if not data or 'markdown' not in data:
            return jsonify({'error': 'No markdown content provided'}), 400
        
        markdown_text = data['markdown']
        filename = data.get('filename', 'document')
        
        if not markdown_text.strip():
            return jsonify({'error': 'Markdown content is empty'}), 400
        
        # Convert markdown to Word document
        doc_buffer = markdown_to_docx(markdown_text, filename)
        
        # Generate filename with timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        word_filename = f"{filename}_{timestamp}.docx"
        
        return send_file(
            doc_buffer,
            as_attachment=True,
            download_name=word_filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        logger.error(f'Error in markdown to word conversion: {str(e)}')
        logger.error(traceback.format_exc())
        return jsonify({'error': f'Conversion error: {str(e)}'}), 500

@app.route('/version', methods=['GET'])
def get_version_info():
    version, release_date, developer = _load_version_meta()
    commit, branch, dirty, describe = _get_git_info()
    return jsonify({
        "version": version,
        "release_date": release_date,
        "developer": developer,
        "git_commit": commit,
        "git_describe": describe,
        "git_branch": branch,
        "git_dirty": dirty,
        "generated_at": datetime.utcnow().isoformat() + "Z",
    })

def convert_docx_to_markdown_sync(docx_path, original_filename):
    """Convert DOCX file to markdown text using Pandoc."""
    try:
        _ensure_pandoc_available()
        logger.debug(f"Attempting to convert DOCX to markdown for filename: {original_filename} using pandoc from path: {docx_path}")
        
        # pypandoc.ensure_pandoc_installed() # Optional, useful for debugging
        
        markdown_output = pypandoc.convert_file(
            docx_path,
            'markdown_strict', # Using markdown_strict for cleaner output, can be 'md' or other flavors
            format='docx'
            # extra_args=['--verbose'] # Uncomment for pandoc verbose logging if needed
        )
        
        logger.info(f"Successfully converted DOCX to markdown for {original_filename} using pandoc.")
        
        # Get file metadata (size, page count is not applicable for DOCX in this context)
        file_size = os.path.getsize(docx_path)

        def format_file_size(size_bytes):
            if size_bytes < 1024:
                return f"{size_bytes} B"
            elif size_bytes < 1024 * 1024:
                return f"{size_bytes / 1024:.1f} KB"
            else:
                return f"{size_bytes / (1024 * 1024):.1f} MB"

        result = {
            'markdown': markdown_output,
            'filename': original_filename,
            'fileSize': format_file_size(file_size),
            'pageCount': None, # Page count is not directly applicable/easy to get for DOCX like PDF
            'timestamp': datetime.now().isoformat(),
            'success': True
        }
        return result
        
    except FileNotFoundError: # Specifically catch if pandoc is not found
        logger.error('Pandoc not found. Please ensure Pandoc is installed and in your PATH.')
        logger.error(traceback.format_exc())
        raise RuntimeError('Pandoc not found. DOCX to Markdown conversion failed.')
    except Exception as e:
        logger.error(f'Error converting DOCX to markdown using Pandoc: {str(e)}')
        logger.error(traceback.format_exc())
        raise e # Re-raise to be handled by the route

@app.route('/convert-word-to-markdown', methods=['POST'])
def convert_word_to_markdown_route():
    temp_path = None
    try:
        if 'document' not in request.files: # Expecting 'document' as the key for docx files
            logger.error('No file in request for Word to Markdown conversion')
            return jsonify({'error': 'No file uploaded'}), 400
            
        file = request.files['document']
        if file.filename == '':
            logger.error('Empty filename for Word to Markdown conversion')
            return jsonify({'error': 'No file selected'}), 400

        if not (file.filename.endswith('.docx')):
            logger.error(f'Invalid file type: {file.filename}. Expected .docx')
            return jsonify({'error': 'Invalid file type. Only .docx files are supported'}), 400

        # Save uploaded file temporarily
        conversion_id = str(uuid.uuid4()) # For unique temp filename
        temp_filename = f'temp_word_upload_{conversion_id}.docx'
        temp_path = os.path.abspath(temp_filename)
        logger.info(f'Saving Word file to {temp_path}')
        file.save(temp_path)
        
        # Perform conversion
        conversion_result = convert_docx_to_markdown_sync(temp_path, file.filename)
        
        return jsonify(conversion_result)
        
    except Exception as e:
        logger.error(f'Server error during Word to Markdown conversion: {str(e)}')
        logger.error(traceback.format_exc())
        return jsonify({'error': f'Server error: {str(e)}', 'success': False}), 500
    finally:
        # Clean up the temporary DOCX file
        if temp_path and os.path.exists(temp_path):
            try:
                os.remove(temp_path)
                logger.info(f"Successfully removed temporary Word upload file: {temp_path}")
            except Exception as e_clean:
                logger.error(f"Error removing temporary Word upload file {temp_path}: {str(e_clean)}")

@app.route("/", defaults={"path": ""})
@app.route("/<path:path>")
def serve_frontend(path):
    """Serve the built frontend (Vite) in production."""
    static_folder = app.static_folder
    if static_folder and os.path.exists(static_folder):
        requested_path = os.path.join(static_folder, path)
        if path and os.path.exists(requested_path):
            return send_from_directory(static_folder, path)
        index_path = os.path.join(static_folder, "index.html")
        if os.path.exists(index_path):
            return send_from_directory(static_folder, "index.html")

    return jsonify({'error': 'Frontend build not found. Run npm run build.'}), 404

if __name__ == '__main__':
    def _free_port(port, retries=5, delay=0.3):
        if os.environ.get("PDF3MD_KILL_PORT", "1") != "1":
            return
        try:
            for attempt in range(retries):
                result = subprocess.run(
                    ["lsof", "-ti", f":{port}"],
                    capture_output=True,
                    text=True,
                    check=False
                )
                pids = [pid.strip() for pid in result.stdout.splitlines() if pid.strip().isdigit()]
                if not pids:
                    return
                for pid in pids:
                    try:
                        os.kill(int(pid), signal.SIGTERM)
                        logger.warning(f"Terminated process on port {port}: PID {pid}")
                    except Exception as kill_error:
                        logger.warning(f"Failed to terminate PID {pid} on port {port}: {kill_error}")
                time.sleep(delay)

            result = subprocess.run(
                ["lsof", "-ti", f":{port}"],
                capture_output=True,
                text=True,
                check=False
            )
            pids = [pid.strip() for pid in result.stdout.splitlines() if pid.strip().isdigit()]
            for pid in pids:
                try:
                    os.kill(int(pid), signal.SIGKILL)
                    logger.warning(f"Force killed process on port {port}: PID {pid}")
                except Exception as kill_error:
                    logger.warning(f"Failed to force kill PID {pid} on port {port}: {kill_error}")
        except FileNotFoundError:
            logger.warning("lsof not found; cannot auto-release port.")
        except Exception as e:
            logger.warning(f"Failed to free port {port}: {e}")

    _free_port(6201)
    logger.info('Starting Flask server...')
    app.run(host='0.0.0.0', port=6201)
