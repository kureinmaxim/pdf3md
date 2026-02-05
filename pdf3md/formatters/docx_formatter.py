"""DOCX document formatting utilities."""

import re
from typing import Dict, Any, Optional
from docx.shared import Pt, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .docx_cleaners import (
    remove_leading_metadata,
    remove_horizontal_rules,
    remove_shape_lines,
)
from .profile_schema import DEFAULT_PROFILE


def _coerce_float(value, default, minimum=None):
    try:
        number = float(value)
    except (TypeError, ValueError):
        number = float(default)
    if minimum is not None and number < minimum:
        return float(minimum)
    return number


def apply_docx_formatting(docx_path: str, profile: Optional[Dict[str, Any]] = None):
    """Apply all formatting to a DOCX document using the specified profile.

    Args:
        docx_path: Path to the DOCX file
        profile: Profile dictionary. If None, uses DEFAULT_PROFILE
    """
    from docx import Document

    if profile is None:
        profile = DEFAULT_PROFILE

    doc = Document(docx_path)

    apply_page_margins(doc, profile)
    remove_leading_metadata(doc)
    try:
        remove_horizontal_rules(doc)
    except FileNotFoundError:
        pass
    remove_shape_lines(doc)
    try:
        add_profile_debug_header(doc, profile)
    except FileNotFoundError:
        pass
    try:
        add_page_numbers(doc, profile)
    except FileNotFoundError:
        pass
    apply_heading_sizes(doc, profile)
    apply_body_font(doc, profile)
    apply_paragraph_formatting(doc, profile)
    format_tables(doc, profile)

    doc.save(docx_path)


def apply_body_font(doc, profile: Dict[str, Any]):
    """Apply body font settings to document.

    Args:
        doc: Document object
        profile: Profile dictionary
    """
    fonts_config = profile.get("fonts", {})
    body_font = fonts_config.get("body", {})
    
    font_name = body_font.get("name", "Calibri")
    font_size = _coerce_float(body_font.get("size", 11), 11, minimum=1)

    # Update Normal style which affects most text
    if 'Normal' in doc.styles:
        style = doc.styles['Normal']
        style.font.name = font_name
        style.font.size = Pt(font_size)
        
        # Explicitly set rFonts to ensure it overrides defaults
        if style.element.rPr is None:
            style.element.get_or_add_rPr()
        
        rFonts = style.element.rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            style.element.rPr.append(rFonts)
            
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:cs'), font_name)

    # Also iterate through all paragraphs to ensure those without explicit style 
    # or with direct formatting overrides get the correct font, IF they are using Normal style
    # or are just plain paragraphs.
    for paragraph in doc.paragraphs:
        # Check if paragraph is using Normal style or no style
        if paragraph.style.name == 'Normal':
            paragraph.style.font.name = font_name
            paragraph.style.font.size = Pt(font_size)
            for run in paragraph.runs:
                # Only override if run doesn't have its own distinct formatting
                # (This is a heuristic, blindly overriding everything might be too aggressive,
                # but for MD conversion usually desired).
                run.font.name = font_name
                run.font.size = Pt(font_size)
                
                # Apply rFonts to runs as well
                if run._element.rPr is None:
                    run._element.get_or_add_rPr()
                rFonts = run._element.rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = OxmlElement('w:rFonts')
                    run._element.rPr.append(rFonts)
                rFonts.set(qn('w:ascii'), font_name)
                rFonts.set(qn('w:hAnsi'), font_name)
                rFonts.set(qn('w:cs'), font_name)


def apply_page_margins(doc, profile: Dict[str, Any]):
    """Apply page margins to document based on profile.

    Args:
        doc: Document object
        profile: Profile dictionary
    """
    page_settings = profile.get("page", {})

    for section in doc.sections:
        section.page_width = Inches(
            _coerce_float(page_settings.get("width", 8.5), 8.5, minimum=0)
        )
        section.page_height = Inches(
            _coerce_float(page_settings.get("height", 11), 11, minimum=0)
        )
        section.top_margin = Inches(
            _coerce_float(page_settings.get("top_margin", 0.3), 0.3, minimum=0)
        )
        section.bottom_margin = Inches(
            _coerce_float(page_settings.get("bottom_margin", 0.3), 0.3, minimum=0)
        )
        section.left_margin = Inches(
            _coerce_float(page_settings.get("left_margin", 0.79), 0.79, minimum=0)
        )
        section.right_margin = Inches(
            _coerce_float(page_settings.get("right_margin", 0.33), 0.33, minimum=0)
        )
        section.header_distance = Inches(
            _coerce_float(page_settings.get("header_distance", 0), 0, minimum=0)
        )
        section.footer_distance = Inches(
            _coerce_float(page_settings.get("footer_distance", 0.2), 0.2, minimum=0)
        )



def add_page_numbers(doc, profile: Dict[str, Any]):
    """Add page numbers to document footer based on profile.

    Args:
        doc: Document object
        profile: Profile dictionary
    """
    page_numbers_config = profile.get("page_numbers", {})

    if not page_numbers_config.get("enabled", True):
        return

    position = page_numbers_config.get("position", "footer_right")
    number_format = page_numbers_config.get("format", "PAGE")
    custom_text = page_numbers_config.get("custom_text", "")
    
    # Determine alignment based on position
    if "right" in position:
        alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif "center" in position:
        alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        alignment = WD_ALIGN_PARAGRAPH.LEFT

    def append_field(paragraph, field_code):
        run = paragraph.add_run()
        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")

        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = f" {field_code} "

        fld_char_separate = OxmlElement("w:fldChar")
        fld_char_separate.set(qn("w:fldCharType"), "separate")

        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")

        run._r.append(fld_char_begin)
        run._r.append(instr_text)
        run._r.append(fld_char_separate)
        run._r.append(fld_char_end)

    for section in doc.sections:
        try:
            footer = section.footer
        except FileNotFoundError:
            continue
        if footer.paragraphs:
            paragraph = footer.paragraphs[0]
        else:
            paragraph = footer.add_paragraph()
        paragraph.alignment = alignment

        if number_format == "PAGE_OF_PAGES":
            append_field(paragraph, "PAGE")
            paragraph.add_run(" of ")
            append_field(paragraph, "NUMPAGES")
        elif number_format == "custom" and custom_text:
            parts = re.split(r"(\{PAGE\}|\{NUMPAGES\})", custom_text)
            for part in parts:
                if part == "{PAGE}":
                    append_field(paragraph, "PAGE")
                elif part == "{NUMPAGES}":
                    append_field(paragraph, "NUMPAGES")
                elif part:
                    paragraph.add_run(part)
        else:
            append_field(paragraph, "PAGE")


def add_profile_debug_header(doc, profile: Dict[str, Any]):
    """Write active profile info into the header for debugging."""
    page = profile.get("page", {})
    fonts = profile.get("fonts", {})
    headings = profile.get("headings", {})
    tables = profile.get("tables", {})
    paragraph = profile.get("paragraph", {})
    page_numbers = profile.get("page_numbers", {})

    header_text = (
        "PROFILE DEBUG: {name} | Page: {w}x{h}in "
        "Margins: T{mt} B{mb} L{ml} R{mr}in | "
        "Body: {body_font} {body_size}pt | "
        "H1: {h1_size}pt {h1_font} | "
        "Tables: min {minw}in max {maxw}in | "
        "Paragraph: line {line} before {before}pt after {after}pt | "
        "PageNumbers: {pn_enabled} {pn_pos} {pn_fmt}"
    ).format(
        name=profile.get("name", "default"),
        w=_coerce_float(page.get("width", 8.5), 8.5, minimum=0),
        h=_coerce_float(page.get("height", 11), 11, minimum=0),
        mt=_coerce_float(page.get("top_margin", 0.3), 0.3, minimum=0),
        mb=_coerce_float(page.get("bottom_margin", 0.3), 0.3, minimum=0),
        ml=_coerce_float(page.get("left_margin", 0.79), 0.79, minimum=0),
        mr=_coerce_float(page.get("right_margin", 0.33), 0.33, minimum=0),
        body_font=fonts.get("body", {}).get("name", "Calibri"),
        body_size=_coerce_float(fonts.get("body", {}).get("size", 11), 11, minimum=1),
        h1_size=_coerce_float(headings.get("h1_size", 14), 14, minimum=1),
        h1_font=fonts.get("heading1", {}).get("name", "Calibri"),
        minw=_coerce_float(tables.get("min_col_width", 0.35), 0.35, minimum=0.1),
        maxw=_coerce_float(tables.get("max_col_width", 3.0), 3.0, minimum=0.1),
        line=_coerce_float(paragraph.get("line_spacing", 1.0), 1.0, minimum=0.1),
        before=_coerce_float(paragraph.get("space_before", 0), 0, minimum=0),
        after=_coerce_float(paragraph.get("space_after", 0), 0, minimum=0),
        pn_enabled=page_numbers.get("enabled", True),
        pn_pos=page_numbers.get("position", "footer_right"),
        pn_fmt=page_numbers.get("format", "PAGE"),
    )

    # Insert at top of document body for visibility
    if doc.paragraphs:
        first_para = doc.paragraphs[0]
        debug_para = first_para.insert_paragraph_before(header_text)
    else:
        debug_para = doc.add_paragraph(header_text)
    for run in debug_para.runs:
        run.font.size = Pt(8)

    # Also write to header (best effort)
    for section in doc.sections:
        try:
            header = section.header
        except FileNotFoundError:
            continue
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        paragraph.text = header_text
        for run in paragraph.runs:
            run.font.size = Pt(8)


def apply_heading_sizes(doc, profile: Dict[str, Any]):
    """Apply font sizes to headings based on profile.

    Args:
        doc: Document object
        profile: Profile dictionary
    """
    headings_config = profile.get("headings", {})
    fonts_config = profile.get("fonts", {})
    body_font_name = fonts_config.get("body", {}).get("name", "Calibri")
    
    heading_sizes = {
        "Heading 1": _coerce_float(headings_config.get("h1_size", 14), 14, minimum=1),
        "Heading 2": _coerce_float(headings_config.get("h2_size", 12), 12, minimum=1),
        "Heading 3": _coerce_float(headings_config.get("h3_size", 11), 11, minimum=1),
        "Heading 4": _coerce_float(headings_config.get("h4_size", 10), 10, minimum=1),
        "Heading 5": _coerce_float(headings_config.get("h5_size", 9), 9, minimum=1),
        "Heading 6": _coerce_float(headings_config.get("h6_size", 9), 9, minimum=1),
    }

    bold = headings_config.get("bold", True)

    heading_fonts = {
        "Heading 1": fonts_config.get("heading1", {}).get("name", body_font_name),
        "Heading 2": fonts_config.get("heading2", {}).get("name", body_font_name),
        "Heading 3": fonts_config.get("heading3", {}).get("name", body_font_name),
        "Heading 4": fonts_config.get("heading4", {}).get("name", body_font_name),
        "Heading 5": fonts_config.get("heading5", {}).get("name", body_font_name),
        "Heading 6": fonts_config.get("heading6", {}).get("name", body_font_name),
    }

    for style_name, size_pt in heading_sizes.items():
        if style_name in doc.styles:
            style = doc.styles[style_name]
            if style and style.font:
                font_name = heading_fonts.get(style_name, body_font_name)
                style.font.size = Pt(size_pt)
                style.font.name = font_name
                if bold:
                    style.font.bold = True
                if style.element.rPr is None:
                    style.element.get_or_add_rPr()
                rFonts = style.element.rPr.find(qn("w:rFonts"))
                if rFonts is None:
                    rFonts = OxmlElement("w:rFonts")
                    style.element.rPr.append(rFonts)
                rFonts.set(qn("w:ascii"), font_name)
                rFonts.set(qn("w:hAnsi"), font_name)
                rFonts.set(qn("w:cs"), font_name)

    for paragraph in doc.paragraphs:
        if paragraph.style and paragraph.style.name in heading_sizes:
            size_pt = heading_sizes[paragraph.style.name]
            font_name = heading_fonts.get(paragraph.style.name, body_font_name)
            for run in paragraph.runs:
                run.font.size = Pt(size_pt)
                run.font.name = font_name
                if bold:
                    run.bold = True
                if run._element.rPr is None:
                    run._element.get_or_add_rPr()
                rFonts = run._element.rPr.find(qn("w:rFonts"))
                if rFonts is None:
                    rFonts = OxmlElement("w:rFonts")
                    run._element.rPr.append(rFonts)
                rFonts.set(qn("w:ascii"), font_name)
                rFonts.set(qn("w:hAnsi"), font_name)
                rFonts.set(qn("w:cs"), font_name)


def apply_paragraph_formatting(doc, profile: Dict[str, Any]):
    """Apply paragraph spacing settings based on profile."""
    paragraph_config = profile.get("paragraph", {})
    line_spacing = _coerce_float(paragraph_config.get("line_spacing", 1.0), 1.0, minimum=0.1)
    space_before = _coerce_float(paragraph_config.get("space_before", 0), 0, minimum=0)
    space_after = _coerce_float(paragraph_config.get("space_after", 0), 0, minimum=0)

    if "Normal" in doc.styles:
        style = doc.styles["Normal"]
        style.paragraph_format.line_spacing = line_spacing
        style.paragraph_format.space_before = Pt(space_before)
        style.paragraph_format.space_after = Pt(space_after)

    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.line_spacing = line_spacing
        paragraph.paragraph_format.space_before = Pt(space_before)
        paragraph.paragraph_format.space_after = Pt(space_after)



def format_tables(doc, profile: Dict[str, Any]):
    """Format all tables in the document based on profile.

    Args:
        doc: Document object
        profile: Profile dictionary
    """
    if not doc.tables:
        return

    tables_config = profile.get("tables", {})
    fonts_config = profile.get("fonts", {})
    
    # Get table font sizes
    table_header_font = fonts_config.get("table_header", {})
    table_body_font = fonts_config.get("table_body", {})
    
    header_font_size = Pt(_coerce_float(table_header_font.get("size", 10), 10, minimum=1))
    body_font_size = Pt(_coerce_float(table_body_font.get("size", 10), 10, minimum=1))
    
    # Table settings
    header_bold = tables_config.get("header_bold", True)
    header_center = tables_config.get("header_center", True)
    min_col_width = Inches(
        _coerce_float(tables_config.get("min_col_width", 0.35), 0.35, minimum=0.1)
    )
    max_col_width = Inches(
        _coerce_float(tables_config.get("max_col_width", 3.0), 3.0, minimum=0.1)
    )
    auto_width = tables_config.get("auto_width", True)

    section = doc.sections[0]
    available_width = section.page_width - section.left_margin - section.right_margin

    for table in doc.tables:
        table.style = "Table"
        table.autofit = not auto_width
        set_table_borders(table, tables_config)

        header_row = table.rows[0] if table.rows else None

        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                set_cell_borders(cell, tables_config)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if header_row and row_index == 0:
                            run.font.size = header_font_size
                            if header_bold:
                                run.bold = True
                        else:
                            run.font.size = body_font_size
                    if header_row and row_index == 0 and header_center:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if header_row and row_index == 0:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        if header_row:
            normalize_header_labels(header_row)

        align_table_columns(table)
        if auto_width:
            adjust_table_column_widths(
                table, available_width, min_col_width, max_col_width, header_row
            )



def adjust_table_column_widths(
    table, available_width, min_col_width, max_col_width, header_row
):
    """Adjust table column widths based on content.

    Args:
        table: Table object
        available_width: Available page width
        min_col_width: Minimum column width
        max_col_width: Maximum column width
        header_row: Header row object
    """
    if not table.columns:
        return

    available_width_emu = int(available_width)
    min_col_width_emu = int(min_col_width)
    max_col_width_emu = int(max_col_width)

    header_texts = []
    if header_row:
        for cell in header_row.cells:
            header_texts.append(get_cell_text(cell))
    else:
        header_texts = [""] * len(table.columns)

    column_max_lengths = []
    per_col_min = []
    for col_index, column in enumerate(table.columns):
        max_len = 1
        for cell in column.cells:
            cell_text = get_cell_text(cell)
            if cell_text:
                max_len = max(max_len, len(cell_text))
        column_max_lengths.append(min(max_len, 120))
        per_col_min.append(min_col_width_emu)

    total_weight = sum(column_max_lengths) or 1
    weighted_lengths = []
    for idx, base in enumerate(column_max_lengths):
        header = header_texts[idx] if idx < len(header_texts) else ""
        weight = float(base)

        header_lower = header.lower()
        header_min_width = get_header_min_width(header)
        if header_min_width:
            per_col_min[idx] = max(per_col_min[idx], int(Inches(header_min_width)))

        if "4-7" in header_lower and "data" in header_lower:
            weight *= 1.2
            if idx < len(per_col_min):
                per_col_min[idx] = int(Inches(0.95))
        elif any(
            key in header_lower
            for key in ["примеч", "назнач", "description", "comment"]
        ):
            weight *= 1.6
        elif any(key in header_lower for key in ["устрой", "device"]):
            weight *= 1.2
        elif any(
            key in header_lower
            for key in ["байт", "byte", "cmd", "code", "op", "id", "№"]
        ):
            weight *= 0.6
        elif "блок" in header_lower:
            weight *= 1.6
            if idx < len(per_col_min):
                per_col_min[idx] = int(Inches(0.9))

        column = table.columns[idx]
        texts = [get_cell_text(cell) for cell in column.cells if get_cell_text(cell)]
        if texts:
            numeric_like = sum(1 for t in texts if is_numeric_like(t))
            if numeric_like / len(texts) >= 0.7:
                weight *= 0.6

        weighted_lengths.append(max(weight, 1))

    total_weight = sum(weighted_lengths) or 1
    raw_widths = [
        available_width_emu * (weight / total_weight) for weight in weighted_lengths
    ]

    clamped_widths = []
    for idx, width in enumerate(raw_widths):
        min_width = per_col_min[idx] if idx < len(per_col_min) else min_col_width_emu
        if width < min_width:
            clamped_widths.append(min_width)
        elif width > max_col_width_emu:
            clamped_widths.append(max_col_width_emu)
        else:
            clamped_widths.append(width)

    total_width = sum(clamped_widths) or available_width_emu
    scale = available_width_emu / total_width if total_width else 1
    final_widths = [int(width * scale) for width in clamped_widths]

    for col_index, width in enumerate(final_widths):
        table.columns[col_index].width = Emu(width)
        for cell in table.columns[col_index].cells:
            cell.width = Emu(width)


def align_table_columns(table):
    """Align table columns based on content type.

    Args:
        table: Table object
    """
    if not table.columns:
        return

    for col_index, column in enumerate(table.columns):
        texts = []
        for cell in column.cells:
            text = get_cell_text(cell)
            if text:
                texts.append(text)

        if not texts:
            continue

        numeric_like = sum(1 for t in texts if is_numeric_like(t))
        avg_len = sum(len(t) for t in texts) / len(texts)
        should_center = (numeric_like / len(texts) >= 0.7) or avg_len <= 4

        if should_center:
            for cell in column.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def is_numeric_like(text):
    """Check if text appears to be numeric.

    Args:
        text: Text to check

    Returns:
        True if text is numeric-like
    """
    value = text.strip()
    if not value:
        return False
    value = value.replace(",", ".")
    if re.fullmatch(r"[+-]?\d+(\.\d+)?", value):
        return True
    if re.fullmatch(r"0x[0-9a-fA-F]+", value):
        return True
    if re.fullmatch(r"[0-9a-fA-F]{2,}", value):
        return True
    return False


def normalize_header_labels(header_row):
    """Normalize and format table header labels.

    Args:
        header_row: Header row object
    """
    for cell in header_row.cells:
        for paragraph in cell.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            if "Блок" in text:
                paragraph.clear()
                run = paragraph.add_run("Блок")
                run.font.size = Pt(10)
                run.bold = True
            if "Устройство" in text:
                paragraph.clear()
                run = paragraph.add_run("Устройство")
                run.font.size = Pt(10)
                run.bold = True
            if "1-й" in text and "байт" in text:
                paragraph.clear()
                run = paragraph.add_run("1-й байт")
                run.font.size = Pt(10)
                run.bold = True
            if "2-й" in text and "байт" in text:
                paragraph.clear()
                run = paragraph.add_run("2-й байт")
                run.font.size = Pt(10)
                run.bold = True
            if "4-7" in text and "Data" in text:
                paragraph.clear()
                run = paragraph.add_run("4-7 ?байт")
                run.font.size = Pt(10)
                run.bold = True
                run.add_break()
                run2 = paragraph.add_run("(Data)")
                run2.font.size = Pt(10)
                run2.bold = True


def get_header_min_width(header_text):
    """Get minimum width for a header based on its content.

    Args:
        header_text: Header text

    Returns:
        Minimum width in inches, or None
    """
    if not header_text:
        return None

    text_lower = header_text.lower()
    if "устрой" in text_lower or "device" in text_lower:
        return 1.2
    if "1-й" in text_lower and "байт" in text_lower:
        return 0.7
    if "2-й" in text_lower and "байт" in text_lower:
        return 0.7
    if "блок" in text_lower:
        return 0.7
    if "4-7" in text_lower and "data" in text_lower:
        return 0.95

    clean_len = len(re.sub(r"\s+", "", header_text))
    if clean_len <= 6:
        return 0.6
    if clean_len <= 12:
        return 0.85
    if clean_len <= 20:
        return 1.2
    return 1.4


def set_table_borders(table, tables_config: Dict[str, Any]):
    """Set borders for a table based on profile.

    Args:
        table: Table object
        tables_config: Tables configuration from profile
    """
    border_style = tables_config.get("border_style", "single")
    border_width = str(tables_config.get("border_width", 8))
    border_color = tables_config.get("border_color", "000000")

    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), border_style)
        element.set(qn("w:sz"), border_width)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), border_color)


def set_cell_borders(cell, tables_config: Dict[str, Any]):
    """Set borders for a table cell based on profile.

    Args:
        cell: Cell object
        tables_config: Tables configuration from profile
    """
    border_style = tables_config.get("border_style", "single")
    border_width = str(tables_config.get("border_width", 8))
    border_color = tables_config.get("border_color", "000000")

    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_tag = qn(f"w:{edge}")
        element = borders.find(edge_tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), border_style)
        element.set(qn("w:sz"), border_width)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), border_color)



def get_cell_text(cell):
    """Extract text from a table cell.

    Args:
        cell: Cell object

    Returns:
        Cell text as string
    """
    parts = []
    for paragraph in cell.paragraphs:
        if paragraph.text:
            parts.append(paragraph.text.strip())
    return " ".join(parts).strip()
