"""DOCX document cleanup utilities."""

from docx.oxml.ns import qn


def delete_paragraph(paragraph):
    """Delete a paragraph from the document.

    Args:
        paragraph: Paragraph object to delete
    """
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def has_horizontal_rule(paragraph):
    """Check if a paragraph contains a horizontal rule.

    Args:
        paragraph: Paragraph to check

    Returns:
        True if paragraph has a horizontal rule
    """
    p_pr = paragraph._element.find(qn("w:pPr"))
    if p_pr is None:
        return False
    p_borders = p_pr.find(qn("w:pBdr"))
    if p_borders is None:
        return False
    for edge in ("top", "bottom", "left", "right"):
        element = p_borders.find(qn(f"w:{edge}"))
        if element is not None and element.get(qn("w:val")):
            return True
    return False


def remove_leading_metadata(doc):
    """Remove leading metadata from the document.

    Args:
        doc: Document object
    """
    paragraphs = list(doc.paragraphs)
    if not paragraphs:
        return

    to_delete = []
    for paragraph in paragraphs[:5]:
        text = paragraph.text.strip()
        if text.startswith("[[") and "#" in text:
            to_delete.append(paragraph)
            continue
        if text in {"---", "***", "___"}:
            to_delete.append(paragraph)
            continue
        if has_horizontal_rule(paragraph):
            to_delete.append(paragraph)

    for paragraph in to_delete:
        delete_paragraph(paragraph)


def remove_horizontal_rules(doc):
    """Remove all horizontal rules from document.

    Args:
        doc: Document object
    """
    for paragraph in list(doc.paragraphs):
        if has_horizontal_rule(paragraph):
            delete_paragraph(paragraph)

    for section in doc.sections:
        try:
            section.header.is_linked_to_previous = False
            section.footer.is_linked_to_previous = False
            for paragraph in list(section.header.paragraphs):
                if has_horizontal_rule(paragraph):
                    delete_paragraph(paragraph)
                else:
                    paragraph.clear()
            for paragraph in list(section.footer.paragraphs):
                if has_horizontal_rule(paragraph):
                    delete_paragraph(paragraph)
        except FileNotFoundError:
            # Some packaged environments miss docx template resources.
            # Skip header/footer cleanup in that case.
            continue


def remove_shape_lines(doc):
    """Remove empty paragraphs containing only shapes/drawings.

    Args:
        doc: Document object
    """
    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip():
            continue
        has_drawing = paragraph._element.findall(
            ".//w:drawing", namespaces=paragraph._element.nsmap
        )
        has_pict = paragraph._element.findall(
            ".//w:pict", namespaces=paragraph._element.nsmap
        )
        if has_drawing or has_pict:
            delete_paragraph(paragraph)
